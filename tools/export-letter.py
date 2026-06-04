#!/usr/bin/env python3
"""
Export a cover letter markdown file to a formatted Word document.

Usage:
    python tools/export-letter.py <letter-name>
    python tools/export-letter.py cover-letter-ncs-group
    python tools/export-letter.py --all

Run from your workspace root. Reads from application-letters/, writes .docx alongside the .md.

Supports two formats:
  - Header format: starts with "Vlad Setchin" name block, date, addressee
  - Dear format:   starts with "Dear <Name>," greeting line
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


LETTERS_DIR = Path.cwd() / "application-letters"

BOLD_LINE_RE = re.compile(r"^\*\*(.*)\*\*$")
HR_RE        = re.compile(r"^---+$")
BLANK_RE     = re.compile(r"^\s*$")
DEAR_RE      = re.compile(r"^Dear\s+.+,$")


def set_run_font(run, size=11, bold=False, name="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name


def add_para(doc, text="", bold=False, size=11,
             space_before=0, space_after=8,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def render_inline(para, text, size=11):
    """Render inline **bold** segments within a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=False)


def detect_format(lines):
    """Return 'dear' if letter starts with Dear <Name>, else 'header'."""
    for line in lines:
        if BLANK_RE.match(line) or HR_RE.match(line) or line.startswith(">"):
            continue
        return "dear" if DEAR_RE.match(line) else "header"
    return "header"


def export_letter(md_path: Path) -> Path:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    fmt   = detect_format(lines)
    doc   = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if fmt == "dear":
        _render_dear(doc, lines)
    else:
        _render_header(doc, lines)

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def _render_dear(doc, lines):
    """Render Dear <Name>, format — no header block."""
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith(">") or HR_RE.match(line):
            i += 1
            continue

        if BLANK_RE.match(line):
            i += 1
            continue

        # Dear <Name>, — greeting
        if DEAR_RE.match(line):
            add_para(doc, line, size=11, space_after=4)
            i += 1
            continue

        # **Re: ...** — subject line
        m = BOLD_LINE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(14)
            run = p.add_run(m.group(1))
            set_run_font(run, size=11, bold=True)
            i += 1
            continue

        # Collect paragraph lines
        para_lines = []
        while i < len(lines) and not BLANK_RE.match(lines[i]) and not HR_RE.match(lines[i]):
            l = lines[i]
            if l.startswith(">") or BOLD_LINE_RE.match(l) or DEAR_RE.match(l):
                break
            para_lines.append(l)
            i += 1

        if not para_lines:
            i += 1
            continue

        text = " ".join(para_lines)
        is_short = len(text) < 80 and "**" not in text
        space_after = 4 if is_short else 10

        # Signature name line — bold
        is_name = (text == "Vlad Setchin")

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)

        if "**" in text:
            render_inline(p, text)
        else:
            run = p.add_run(text)
            set_run_font(run, size=11, bold=is_name)


def _render_header(doc, lines):
    """Render original header format — name block, date, addressee, Re:."""
    i = 0
    first_name_done = False

    while i < len(lines):
        line = lines[i]

        if line.startswith(">") or HR_RE.match(line):
            i += 1
            continue

        if BLANK_RE.match(line):
            i += 1
            continue

        # First non-blank line = name (large, bold)
        if not first_name_done:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            set_run_font(run, size=14, bold=True)
            first_name_done = True
            i += 1
            continue

        # Bold-only line — Re: or section header
        m = BOLD_LINE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(14)
            run = p.add_run(m.group(1))
            set_run_font(run, size=11, bold=True)
            i += 1
            continue

        # Collect paragraph
        para_lines = []
        while i < len(lines) and not BLANK_RE.match(lines[i]) and not HR_RE.match(lines[i]):
            l = lines[i]
            if l.startswith(">") or BOLD_LINE_RE.match(l):
                break
            para_lines.append(l)
            i += 1

        if not para_lines:
            i += 1
            continue

        text = " ".join(para_lines)
        is_short    = len(text) < 80 and "**" not in text
        space_after = 4 if is_short else 10

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)

        if "**" in text:
            render_inline(p, text)
        else:
            run = p.add_run(text)
            set_run_font(run, size=11)


def main():
    args = sys.argv[1:]

    if not args:
        print(__doc__)
        sys.exit(1)

    if args[0] == "--all":
        md_files = list(LETTERS_DIR.glob("cover-letter-*.md"))
        if not md_files:
            print(f"No cover letter .md files found in {LETTERS_DIR}")
            sys.exit(1)
    else:
        name = args[0]
        if not name.endswith(".md"):
            name = name + ".md"
        md_path = LETTERS_DIR / name
        if not md_path.exists():
            print(f"ERROR: File not found: {md_path}")
            print("Available letters:")
            for f in sorted(LETTERS_DIR.glob("cover-letter-*.md")):
                print(f"  {f.stem}")
            sys.exit(1)
        md_files = [md_path]

    for md_path in md_files:
        out = export_letter(md_path)
        print(f"Exported: {out}")


if __name__ == "__main__":
    main()
